#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成模块
"""

import os
import json
from datetime import datetime
from jinja2 import Template
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGenerator:
    """报告生成器类"""
    
    def __init__(self, results, target, output_dir):
        """初始化报告生成器
        
        Args:
            results: 扫描结果
            target: 扫描目标
            output_dir: 输出目录
        """
        self.results = results
        self.target = target
        self.output_dir = output_dir
        self.timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    def generate_html(self, filename='security_report.html', template_file=None):
        """生成HTML报告"""
        # 尝试从模板文件加载
        if template_file and os.path.exists(template_file):
            with open(template_file, 'r', encoding='utf-8') as f:
                template_content = f.read()
            template = Template(template_content)
        # 尝试从templates目录加载（检查多个可能的位置）
        else:
            # 计算相对于当前文件的模板目录路径
            script_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(script_dir)
            template_dir = os.path.join(project_root, 'templates')
            
            # 检查templates目录是否存在
            if os.path.exists(template_dir):
                template_path = os.path.join(template_dir, 'html_template.html')
                if os.path.exists(template_path):
                    with open(template_path, 'r', encoding='utf-8') as f:
                        template_content = f.read()
                    template = Template(template_content)
                else:
                    # 使用默认模板
                    template = Template('''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>HOS-LS 安全检测报告</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css">
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            background-color: #f5f7fa;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background-color: white;
            padding: 30px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            min-height: 100vh;
        }
        h1 {
            color: #2c3e50;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 20px;
            margin-bottom: 30px;
            font-size: 2.5em;
        }
        h2 {
            color: #34495e;
            margin-top: 40px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 10px;
            font-size: 1.8em;
        }
        h3 {
            color: #7f8c8d;
            margin-top: 20px;
            font-size: 1.4em;
        }
        .summary {
            background-color: #f9f9f9;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        .summary-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }
        .summary-item {
            text-align: center;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        .summary-item h3 {
            margin: 0 0 10px 0;
            font-size: 1.2em;
            color: #7f8c8d;
        }
        .summary-item .value {
            font-size: 2em;
            font-weight: bold;
        }
        .high-risk {
            background-color: #ffebee;
            color: #c62828;
        }
        .medium-risk {
            background-color: #fff3e0;
            color: #ef6c00;
        }
        .low-risk {
            background-color: #e8f5e8;
            color: #2e7d32;
        }
        .risk-level {
            padding: 8px 16px;
            border-radius: 20px;
            font-weight: bold;
            display: inline-block;
        }
        .risk-high {
            background-color: #ffebee;
            color: #c62828;
        }
        .risk-medium {
            background-color: #fff3e0;
            color: #ef6c00;
        }
        .risk-low {
            background-color: #e8f5e8;
            color: #2e7d32;
        }
        .project-info {
            background-color: #e3f2fd;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        .project-info-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 15px;
        }
        .project-info-item {
            display: flex;
            align-items: center;
        }
        .project-info-item i {
            margin-right: 10px;
            color: #3498db;
        }
        .security-assessment {
            background-color: #f3e5f5;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        .recommendations {
            background-color: #fff8e1;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        .recommendation-item {
            margin-bottom: 15px;
            padding: 15px;
            background-color: white;
            border-radius: 6px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        }
        .recommendation-item h4 {
            margin: 0 0 10px 0;
            color: #34495e;
        }
        .recommendation-item p {
            margin: 0;
            color: #7f8c8d;
        }
        .chart-container {
            position: relative;
            height: 400px;
            margin: 30px 0;
        }
        .tabs {
            margin: 30px 0;
        }
        .tab-buttons {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin-bottom: 20px;
        }
        .tab-button {
            background-color: #f1f1f1;
            border: 1px solid #ddd;
            padding: 12px 20px;
            cursor: pointer;
            border-radius: 4px;
            transition: all 0.3s ease;
            font-size: 14px;
        }
        .tab-button:hover {
            background-color: #e0e0e0;
        }
        .tab-button.active {
            background-color: #3498db;
            color: white;
            border-color: #3498db;
        }
        .tab-content {
            border: 1px solid #ddd;
            border-radius: 4px;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .filter-container {
            margin-bottom: 20px;
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
        }
        .filter-input {
            padding: 10px;
            border: 1px solid #ddd;
            border-radius: 4px;
            flex: 1;
            min-width: 200px;
        }
        .filter-select {
            padding: 10px;
            border: 1px solid #ddd;
            border-radius: 4px;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            background-color: white;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
            color: #34495e;
            position: sticky;
            top: 0;
            z-index: 10;
        }
        tr:hover {
            background-color: #f5f5f5;
        }
        tr.highlight {
            background-color: #fff3cd;
        }
        .footer {
            margin-top: 60px;
            text-align: center;
            color: #666;
            font-size: 14px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }
        .badge {
            padding: 4px 8px;
            border-radius: 12px;
            font-size: 12px;
            font-weight: bold;
        }
        .badge-high {
            background-color: #dc3545;
            color: white;
        }
        .badge-medium {
            background-color: #ffc107;
            color: #212529;
        }
        .badge-low {
            background-color: #28a745;
            color: white;
        }
        .code-snippet {
            background-color: #f8f9fa;
            padding: 10px;
            border-radius: 4px;
            font-family: 'Courier New', Courier, monospace;
            font-size: 13px;
            margin: 10px 0;
            white-space: pre-wrap;
            border-left: 3px solid #3498db;
        }
        .search-container {
            margin-bottom: 20px;
        }
        .search-input {
            width: 100%;
            padding: 12px;
            border: 1px solid #ddd;
            border-radius: 4px;
            font-size: 14px;
        }
        .priority-high {
            border-left: 4px solid #dc3545;
        }
        .priority-medium {
            border-left: 4px solid #ffc107;
        }
        .priority-low {
            border-left: 4px solid #28a745;
        }
        @media (max-width: 768px) {
            .container {
                padding: 20px;
            }
            .summary-grid {
                grid-template-columns: 1fr;
            }
            .project-info-grid {
                grid-template-columns: 1fr;
            }
            .tab-buttons {
                flex-direction: column;
            }
            .tab-button {
                width: 100%;
                text-align: left;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1><i class="fas fa-shield-alt"></i> HOS-LS 安全检测报告</h1>
        
        <div class="project-info">
            <h2>项目信息</h2>
            <div class="project-info-grid">
                <div class="project-info-item">
                    <i class="fas fa-folder-open fa-lg"></i>
                    <div>
                        <strong>检测目标:</strong> {{ target }}
                    </div>
                </div>
                <div class="project-info-item">
                    <i class="fas fa-clock fa-lg"></i>
                    <div>
                        <strong>检测时间:</strong> {{ timestamp }}
                    </div>
                </div>
                <div class="project-info-item">
                    <i class="fas fa-code fa-lg"></i>
                    <div>
                        <strong>项目类型:</strong> {{ project_type }}
                    </div>
                </div>
                <div class="project-info-item">
                    <i class="fas fa-list-check fa-lg"></i>
                    <div>
                        <strong>使用规则集:</strong> {{ rule_set }}
                    </div>
                </div>
            </div>
        </div>
        
        <div class="summary">
            <h2>检测摘要</h2>
            <div class="summary-grid">
                <div class="summary-item high-risk">
                    <h3><i class="fas fa-exclamation-triangle"></i> 高风险</h3>
                    <div class="value">{{ high_risk }}</div>
                </div>
                <div class="summary-item medium-risk">
                    <h3><i class="fas fa-exclamation-circle"></i> 中风险</h3>
                    <div class="value">{{ medium_risk }}</div>
                </div>
                <div class="summary-item low-risk">
                    <h3><i class="fas fa-info-circle"></i> 低风险</h3>
                    <div class="value">{{ low_risk }}</div>
                </div>
                <div class="summary-item" style="background-color: #e3f2fd; color: #1976d2;">
                    <h3><i class="fas fa-shield-alt"></i> 总体风险</h3>
                    <div class="value">{{ overall_risk_text }}</div>
                </div>
            </div>
            
            <div class="chart-container">
                <canvas id="riskChart"></canvas>
            </div>
        </div>
        
        <div class="security-assessment">
            <h2><i class="fas fa-search"></i> 安全评估</h2>
            <p>{{ security_assessment }}</p>
        </div>
        
        <div class="recommendations">
            <h2><i class="fas fa-lightbulb"></i> 安全建议</h2>
            {% for recommendation in recommendations %}
            <div class="recommendation-item">
                <h4><span class="badge badge-{{ recommendation.severity }}">{{ recommendation.severity }}风险</span> {{ recommendation.text }}</h4>
                {% if recommendation.details %}
                <p>{{ recommendation.details }}</p>
                {% endif %}
            </div>
            {% endfor %}
        </div>
        
        <div class="tabs">
            <h2><i class="fas fa-list"></i> 详细检测结果</h2>
            <div class="tab-buttons">
                <button class="tab-button active" data-tab="all">全部问题</button>
                <button class="tab-button" data-tab="code_security">代码安全</button>
                <button class="tab-button" data-tab="ai_security">AI安全</button>
                <button class="tab-button" data-tab="permission_security">权限安全</button>
                <button class="tab-button" data-tab="network_security">网络安全</button>
                <button class="tab-button" data-tab="dependency_security">依赖安全</button>
            </div>
            
            <div class="tab-content">
                <div class="search-container">
                    <input type="text" class="search-input" placeholder="搜索问题..." id="searchInput">
                </div>
                
                <div class="filter-container">
                    <select class="filter-select" id="severityFilter">
                        <option value="all">所有严重程度</option>
                        <option value="high">高风险</option>
                        <option value="medium">中风险</option>
                        <option value="low">低风险</option>
                    </select>
                    <select class="filter-select" id="categoryFilter">
                        <option value="all">所有类别</option>
                        <option value="code_security">代码安全</option>
                        <option value="ai_security">AI安全</option>
                        <option value="permission_security">权限安全</option>
                        <option value="network_security">网络安全</option>
                        <option value="dependency_security">依赖安全</option>
                    </select>
                </div>
                
                <table id="resultsTable">
                    <thead>
                        <tr>
                            <th>文件</th>
                            <th>问题</th>
                            <th>严重程度</th>
                            <th>类别</th>
                            <th>详情</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for category, items in all_issues.items() %}
                            {% for item in items %}
                            <tr class="priority-{{ item.severity }}" data-category="{{ category }}" data-severity="{{ item.severity }}">
                                <td>{{ item.file }}</td>
                                <td>{{ item.issue }}</td>
                                <td><span class="badge badge-{{ item.severity }}">{{ item.severity }}</span></td>
                                <td>{{ category }}</td>
                                <td>
                                    <div>{{ item.details }}</div>
                                    {% if item.code_snippet %}
                                    <div class="code-snippet">{{ item.code_snippet }}</div>
                                    {% endif %}
                                </td>
                            </tr>
                            {% endfor %}
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        </div>
        
        {% if ai_suggestions %}
        <div class="ai-suggestions">
            <h2><i class="fas fa-robot"></i> AI 辅助建议</h2>
            <div class="recommendations">
                {% if ai_suggestions.risk_assessment %}
                <h3>风险评估</h3>
                <p>{{ ai_suggestions.risk_assessment }}</p>
                {% endif %}
                
                {% if ai_suggestions.specific_suggestions %}
                <h3>具体建议</h3>
                {% for suggestion in ai_suggestions.specific_suggestions %}
                <div class="recommendation-item">
                    <h4>{{ suggestion }}</h4>
                </div>
                {% endfor %}
                {% endif %}
                
                {% if ai_suggestions.best_practices %}
                <h3>最佳实践</h3>
                {% for practice in ai_suggestions.best_practices %}
                <div class="recommendation-item">
                    <h4>{{ practice }}</h4>
                </div>
                {% endfor %}
                {% endif %}
            </div>
        </div>
        {% endif %}
        
        <div class="footer">
            <p>报告生成时间: {{ timestamp }}</p>
            <p>HOS-LS 安全检测工具 v1.0</p>
        </div>
    </div>
    
    <script>
        // 风险分布图表
        const ctx = document.getElementById('riskChart').getContext('2d');
        const riskChart = new Chart(ctx, {
            type: 'pie',
            data: {
                labels: ['高风险', '中风险', '低风险'],
                datasets: [{
                    data: [{{ high_risk }}, {{ medium_risk }}, {{ low_risk }}],
                    backgroundColor: [
                        '#dc3545',
                        '#ffc107',
                        '#28a745'
                    ],
                    borderColor: [
                        '#c82333',
                        '#e0a800',
                        '#218838'
                    ],
                    borderWidth: 1
                }]
            },
            options: {
                responsive: true,
                plugins: {
                    legend: {
                        position: 'bottom',
                    },
                    title: {
                        display: true,
                        text: '风险分布'
                    }
                }
            }
        });
        
        // 选项卡功能
        document.querySelectorAll('.tab-button').forEach(button => {
            button.addEventListener('click', () => {
                // 移除所有活动状态
                document.querySelectorAll('.tab-button').forEach(btn => btn.classList.remove('active'));
                // 添加当前按钮活动状态
                button.classList.add('active');
                
                const tab = button.getAttribute('data-tab');
                filterResults(tab);
            });
        });
        
        // 搜索功能
        document.getElementById('searchInput').addEventListener('input', () => {
            filterResults(document.querySelector('.tab-button.active').getAttribute('data-tab'));
        });
        
        // 过滤功能
        document.getElementById('severityFilter').addEventListener('change', () => {
            filterResults(document.querySelector('.tab-button.active').getAttribute('data-tab'));
        });
        
        document.getElementById('categoryFilter').addEventListener('change', () => {
            filterResults(document.querySelector('.tab-button.active').getAttribute('data-tab'));
        });
        
        function filterResults(tab) {
            const searchTerm = document.getElementById('searchInput').value.toLowerCase();
            const severityFilter = document.getElementById('severityFilter').value;
            const categoryFilter = document.getElementById('categoryFilter').value;
            
            const rows = document.querySelectorAll('#resultsTable tbody tr');
            rows.forEach(row => {
                const category = row.getAttribute('data-category');
                const severity = row.getAttribute('data-severity');
                const text = row.textContent.toLowerCase();
                
                const matchesTab = tab === 'all' || category === tab;
                const matchesSearch = text.includes(searchTerm);
                const matchesSeverity = severityFilter === 'all' || severity === severityFilter;
                const matchesCategory = categoryFilter === 'all' || category === categoryFilter;
                
                if (matchesTab && matchesSearch && matchesSeverity && matchesCategory) {
                    row.style.display = '';
                } else {
                    row.style.display = 'none';
                }
            });
        }
        
        // 初始过滤
        filterResults('all');
    </script>
</body>
</html>
''')
        
        # 计算风险数量
        high_risk = 0
        medium_risk = 0
        low_risk = 0
        
        for category in self.results.values():
            if isinstance(category, list):
                for item in category:
                    if item.get('severity') == 'high':
                        high_risk += 1
                    elif item.get('severity') == 'medium':
                        medium_risk += 1
                    elif item.get('severity') == 'low':
                        low_risk += 1
        
        # 计算总体风险等级
        if high_risk > 0:
            overall_risk = 'high'
            overall_risk_text = '高风险'
        elif medium_risk > 3:
            overall_risk = 'medium'
            overall_risk_text = '中风险'
        else:
            overall_risk = 'low'
            overall_risk_text = '低风险'
        
        # 获取项目类型和规则集
        project_type = self.results.get('project_type', 'unknown')
        rule_set = self.results.get('rule_set', 'default')
        
        # 生成安全评估
        security_assessment = f"本次安全扫描共发现 {high_risk} 个高风险问题，{medium_risk} 个中风险问题，{low_risk} 个低风险问题。"
        if high_risk > 0:
            security_assessment += " 系统存在严重安全隐患，建议立即处理高风险问题。"
        elif medium_risk > 3:
            security_assessment += " 系统存在一定安全风险，建议尽快处理中风险问题。"
        else:
            security_assessment += " 系统安全状态良好，建议定期进行安全检查。"
        
        # 生成安全建议
        recommendations = []
        if high_risk > 0:
            recommendations.append({"severity": "高", "text": "立即处理所有高风险问题，特别是硬编码的敏感信息和后门代码。"})
        if medium_risk > 0:
            recommendations.append({"severity": "中", "text": "尽快处理中风险问题，如网络访问代码的安全验证和依赖库版本管理。"})
        if low_risk > 0:
            recommendations.append({"severity": "低", "text": "定期进行安全检查，保持系统和依赖库的更新。"})
        
        # 根据项目类型生成特定建议
        if project_type == "openclaw":
            recommendations.append({"severity": "中", "text": "确保OpenClaw端口仅本地监听，避免公网暴露。"})
            recommendations.append({"severity": "中", "text": "定期更新OpenClaw到最新版本，获取安全修复。"})
        elif project_type == "cursor":
            recommendations.append({"severity": "中", "text": "确保Cursor生成的代码不包含硬编码的API密钥。"})
            recommendations.append({"severity": "低", "text": "审查Cursor生成的提示模板，确保不包含敏感信息。"})
        
        # 通用安全建议
        recommendations.append({"severity": "低", "text": "使用环境变量存储敏感信息，避免硬编码。"})
        recommendations.append({"severity": "低", "text": "遵循最小权限原则，限制文件和目录权限。"})
        recommendations.append({"severity": "低", "text": "定期进行安全扫描，及时发现和处理安全问题。"})
        
        # 获取AI建议
        ai_suggestions = self.results.get('ai_suggestions', {})
        
        html = template.render(
            target=self.target,
            timestamp=self.timestamp,
            high_risk=high_risk,
            medium_risk=medium_risk,
            low_risk=low_risk,
            overall_risk=overall_risk,
            overall_risk_text=overall_risk_text,
            security_assessment=security_assessment,
            recommendations=recommendations,
            project_type=project_type,
            rule_set=rule_set,
            code_security=self.results.get('code_security', []),
            permission_security=self.results.get('permission_security', []),
            network_security=self.results.get('network_security', []),
            dependency_security=self.results.get('dependency_security', []),
            config_security=self.results.get('config_security', []),
            ai_suggestions=ai_suggestions
        )
        
        output_path = os.path.join(self.output_dir, filename)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return output_path
    
    def generate_docx(self, filename='security_report.docx'):
        """生成DOCX报告"""
        doc = Document()
        
        # 添加标题
        doc.add_heading('HOS-LS 安全检测报告', 0)
        
        # 添加检测摘要
        doc.add_heading('检测摘要', level=1)
        
        # 计算风险数量
        high_risk = 0
        medium_risk = 0
        low_risk = 0
        
        for category in self.results.values():
            if isinstance(category, list):
                for item in category:
                    if item['severity'] == 'high':
                        high_risk += 1
                    elif item['severity'] == 'medium':
                        medium_risk += 1
                    elif item['severity'] == 'low':
                        low_risk += 1
        
        # 添加摘要内容
        summary_table = doc.add_table(rows=1, cols=5)
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = '检测目标'
        hdr_cells[1].text = '检测时间'
        hdr_cells[2].text = '高风险'
        hdr_cells[3].text = '中风险'
        hdr_cells[4].text = '低风险'
        
        row_cells = summary_table.add_row().cells
        row_cells[0].text = self.target
        row_cells[1].text = self.timestamp
        row_cells[2].text = str(high_risk)
        row_cells[3].text = str(medium_risk)
        row_cells[4].text = str(low_risk)
        
        # 添加详细内容
        categories = [
            ('代码安全', 'code_security'),
            ('权限安全', 'permission_security'),
            ('网络安全', 'network_security'),
            ('依赖安全', 'dependency_security'),
            ('配置安全', 'config_security')
        ]
        
        for category_name, category_key in categories:
            items = self.results.get(category_key, [])
            if items:
                doc.add_heading(category_name, level=1)
                table = doc.add_table(rows=1, cols=5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '文件'
                hdr_cells[1].text = '行号'
                hdr_cells[2].text = '问题'
                hdr_cells[3].text = '严重程度'
                hdr_cells[4].text = '详情'
                
                for item in items:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item['file']
                    row_cells[1].text = str(item.get('line_number', ''))
                    row_cells[2].text = item['issue']
                    row_cells[3].text = item['severity']
                    row_cells[4].text = item['details']
        
        # 添加页脚
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.add_run(f'报告生成时间: {self.timestamp}')
        footer_paragraph.add_run('\nHOS-LS 安全检测工具 v1.0.0')
        
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        
        return output_path
    
    def generate_md(self, filename='security_report.md', template_file=None):
        """生成MD报告"""
        # 计算风险数量
        high_risk = 0
        medium_risk = 0
        low_risk = 0
        
        for category in self.results.values():
            if isinstance(category, list):
                for item in category:
                    if item['severity'] == 'high':
                        high_risk += 1
                    elif item['severity'] == 'medium':
                        medium_risk += 1
                    elif item['severity'] == 'low':
                        low_risk += 1
        
        # 尝试从模板文件加载
        if template_file and os.path.exists(template_file):
            with open(template_file, 'r', encoding='utf-8') as f:
                template_content = f.read()
            template = Template(template_content)
            md = template.render(
                target=self.target,
                timestamp=self.timestamp,
                high_risk=high_risk,
                medium_risk=medium_risk,
                low_risk=low_risk,
                code_security=self.results.get('code_security', []),
                permission_security=self.results.get('permission_security', []),
                network_security=self.results.get('network_security', []),
                dependency_security=self.results.get('dependency_security', []),
                config_security=self.results.get('config_security', [])
            )
        # 尝试从templates目录加载
        else:
            # 计算相对于当前文件的模板目录路径
            script_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(script_dir)
            template_dir = os.path.join(project_root, 'templates')
            
            # 检查templates目录是否存在
            if os.path.exists(template_dir):
                template_path = os.path.join(template_dir, 'md_template.md')
                if os.path.exists(template_path):
                    with open(template_path, 'r', encoding='utf-8') as f:
                        template_content = f.read()
                    template = Template(template_content)
                    md = template.render(
                        target=self.target,
                        timestamp=self.timestamp,
                        high_risk=high_risk,
                        medium_risk=medium_risk,
                        low_risk=low_risk,
                        code_security=self.results.get('code_security', []),
                        permission_security=self.results.get('permission_security', []),
                        network_security=self.results.get('network_security', []),
                        dependency_security=self.results.get('dependency_security', []),
                        config_security=self.results.get('config_security', [])
                    )
            else:
                # 使用默认模板
                md = f'''
# HOS-LS 安全检测报告

## 检测摘要

| 检测目标 | 检测时间 | 高风险 | 中风险 | 低风险 |
|---------|---------|-------|-------|-------|
| {self.target} | {self.timestamp} | {high_risk} | {medium_risk} | {low_risk} |
'''
                
                # 添加详细内容
                categories = [
                    ('代码安全', 'code_security'),
                    ('权限安全', 'permission_security'),
                    ('网络安全', 'network_security'),
                    ('依赖安全', 'dependency_security'),
                    ('配置安全', 'config_security')
                ]
                
                for category_name, category_key in categories:
                    items = self.results.get(category_key, [])
                    if items:
                        md += f'''
## {category_name}

| 文件 | 问题 | 严重程度 | 详情 |
|------|------|---------|------|
'''
                        for item in items:
                            md += f'| {item["file"]} | {item["issue"]} | {item["severity"]} | {item["details"]} |\n'
                
                md += f'''
---

报告生成时间: {self.timestamp}
HOS-LS 安全检测工具 v1.0.0
'''
        
        output_path = os.path.join(self.output_dir, filename)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md)
        
        return output_path
